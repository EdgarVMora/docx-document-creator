# Importaciones necesarias para la generación de documentos
from docx import Document  # Para crear y manipular documentos de Word
from docx.shared import Pt, Inches, RGBColor  # Utilidades para formato de documento
import os  # Para manejo de rutas de archivos

def generar_documento(data, carpeta_imagenes, carpeta_salida):
    """
    Genera un documento Word con la información del árbol proporcionada.
    
    Args:
        data (dict): Diccionario con los datos del árbol
        carpeta_imagenes (str): Ruta a la carpeta donde se encuentran las imágenes
        carpeta_salida (str): Ruta donde se guardará el documento generado
    """
    # Crear un nuevo documento
    doc = Document()

    # Agregar título principal con emoji de árbol
    doc.add_heading(f"🌳 {data['nombre']}", level=0)
    doc.add_paragraph("")  # Espacio después del título

    # Agregar descripción con formato destacado
    parrafo_desc = doc.add_paragraph()
    run_desc = parrafo_desc.add_run(data['descripcion'])
    run_desc.bold = True  # Texto en negrita
    run_desc.font.size = Pt(12)  # Tamaño de fuente
    run_desc.font.color.rgb = RGBColor(0, 102, 204)  # Color azul
    doc.add_paragraph("")  # Espacio después de la descripción

    # Crear tabla con información técnica básica
    doc.add_paragraph("📊 Información Técnica:", style='Intense Quote')
    tabla = doc.add_table(rows=5, cols=2)
    tabla.style = 'Light Grid Accent 1'

    # Definir campos de la tabla técnica con emojis
    campos = [
        ("📍 Ubicación", data.get("ubicacion", "No especificada")),
        ("🧬 Especie", data.get("especie", "Desconocida")),
        ("📏 Altura (m)", str(data.get("altura_metros", "N/A"))),
        ("⏳ Edad Aproximada", data.get("edad_aproximada", "N/A")),
        ("❤️ Estado de Salud", data.get("estado_salud", "No registrado"))
    ]

    # Llenar la tabla con los campos definidos
    for i, (campo, valor) in enumerate(campos):
        tabla.cell(i, 0).text = campo
        tabla.cell(i, 1).text = valor

    doc.add_paragraph("")  # Espacio después de la tabla

    # Agregar pie de tabla si existe
    if "pie_tabla" in data:
        pie_tabla = doc.add_paragraph()
        run_pie_tabla = pie_tabla.add_run(data["pie_tabla"])
        run_pie_tabla.italic = True  # Texto en cursiva
        run_pie_tabla.font.size = Pt(9)  # Tamaño pequeño
        run_pie_tabla.font.color.rgb = RGBColor(100, 100, 100)  # Color gris
        doc.add_paragraph("")

    # Agregar tabla extendida si existe información adicional
    if "tabla_extendida" in data:
        doc.add_paragraph("📑 Información Extendida:")
        tabla_ext = doc.add_table(rows=1, cols=2)
        tabla_ext.style = 'Table Grid'
        tabla_ext.autofit = True

        # Configurar encabezados de la tabla extendida
        tabla_ext.cell(0, 0).text = "Atributo"
        tabla_ext.cell(0, 1).text = "Valor"

        # Llenar la tabla extendida con los datos proporcionados
        for entrada in data["tabla_extendida"]:
            fila = tabla_ext.add_row()
            fila.cells[0].text = entrada["atributo"]
            fila.cells[1].text = entrada["valor"]

        doc.add_paragraph("")  # Espacio después de la tabla extendida

    # Agregar fecha del registro
    doc.add_paragraph(f"📅 Fecha de registro: {data['fecha']}")
    doc.add_paragraph("")

    # Procesar y agregar imagen si existe
    imagen_path = os.path.join(carpeta_imagenes, data.get("imagen", ""))
    if os.path.exists(imagen_path):
        doc.add_paragraph("📸 Fotografía del Árbol:")
        doc.add_picture(imagen_path, width=Inches(4.0))  # Ancho fijo de 4 pulgadas
        
        # Agregar pie de imagen si existe
        if "pie_imagen" in data:
            pie = doc.add_paragraph()
            run_pie = pie.add_run(data["pie_imagen"])
            run_pie.italic = True  # Texto en cursiva
            run_pie.font.size = Pt(9)  # Tamaño pequeño
            run_pie.font.color.rgb = RGBColor(100, 100, 100)  # Color gris
    else:
        doc.add_paragraph("⚠️ Imagen no disponible.")

    # Generar nombre del archivo y guardar el documento
    filename = f"{data['id']}_{data['nombre'].replace(' ', '_')}.docx"
    ruta_salida = os.path.join(carpeta_salida, filename)
    doc.save(ruta_salida)

    print(f"✅ Documento generado: {ruta_salida}")